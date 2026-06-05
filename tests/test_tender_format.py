from __future__ import annotations

from docx import Document

from bid_document_service.tender_format import (
    add_response_front_matter,
    amount_to_rmb_upper,
    cleanup_empty_headings,
    element_text,
    fill_fields,
    fill_table_rows,
    find_format_range,
    insert_sections,
    normalize_money_text,
)


def test_find_format_range_skips_table_of_contents_entry():
    doc = Document()
    doc.add_paragraph("目 录")
    doc.add_paragraph("第三部分 响应文件格式\t18")
    doc.add_paragraph("第四部分 采购项目内容及要求\t31")
    doc.add_paragraph("第一部分 磋商邀请")
    doc.add_paragraph("第二部分 参与磋商供应商须知")
    doc.add_paragraph("第三部分 响应文件格式")
    doc.add_paragraph("一、报价表")
    doc.add_paragraph("第四部分 采购项目内容及要求")

    start_idx, end_idx, warnings = find_format_range(doc)
    children = list(doc._body._element)

    assert start_idx > 2
    assert element_text(children[start_idx]) == "第三部分 响应文件格式"
    assert element_text(children[end_idx]) == "第四部分 采购项目内容及要求"
    assert any("已跳过目录页" in warning for warning in warnings)


def test_find_format_range_skips_explanatory_format_requirement():
    doc = Document()
    doc.add_paragraph("第二部分 参与磋商供应商须知")
    doc.add_paragraph("（3）响应文件格式要求；")
    doc.add_paragraph("（4）供应商有关资格证明文件要求；")
    doc.add_paragraph("供应商应认真阅读采购文件中所有事项、格式条款和规范要求。")
    doc.add_paragraph("13．响应文件组成")
    doc.add_paragraph("响应供应商应按照采购文件的规定和要求编制响应文件。")
    doc.add_paragraph("第三部分 响应文件格式")
    doc.add_paragraph("一、报价表")
    doc.add_paragraph("第四部分 采购项目内容及要求")

    start_idx, end_idx, warnings = find_format_range(doc)
    children = list(doc._body._element)

    assert element_text(children[start_idx]) == "第三部分 响应文件格式"
    assert element_text(children[end_idx]) == "第四部分 采购项目内容及要求"
    assert any("已跳过说明性条款" in warning for warning in warnings)
    assert any("已跳过供应商须知中的响应文件组成说明" in warning for warning in warnings)


def test_insert_sections_expands_json_payload_under_service_heading():
    doc = Document()
    doc.add_paragraph("第三部分 响应文件格式")
    doc.add_paragraph("十、服务方案")
    doc.add_paragraph("十一、其它应答")
    payload = """
    {
      "推荐响应文件目录": [{"章节编号": "1", "标题": "磋商函"}],
      "方案章节": [
        {
          "章节编号": "10.1.1",
          "层级": 3,
          "标题": "项目概述与需求理解",
          "正文": "围绕统一加密平台升级目标进行响应。",
          "编制要点": ["覆盖实质性要求", "列明证明材料"]
        }
      ]
    }
    """

    warnings = insert_sections(doc, [{"章节编号": "10", "层级": 1, "标题": "服务方案", "正文": payload}])
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

    assert '"方案章节"' not in text
    assert "10.1.1 项目概述与需求理解" in text
    assert "围绕统一加密平台升级目标进行响应。" in text
    assert any("已插入到招标原格式的“服务方案/技术方案”章节下" in warning for warning in warnings)


def test_fill_fields_populates_bidder_aliases_in_text_and_table_cells():
    doc = Document()
    doc.add_paragraph("参与磋商供应商名称：（盖章）")
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "供应商名称"

    warnings = fill_fields(doc, {"供应商名称": "示例科技有限公司"})

    assert "参与磋商供应商名称：示例科技有限公司" in doc.paragraphs[0].text
    assert table.rows[0].cells[1].text == "示例科技有限公司"
    assert not warnings


def test_front_matter_removes_format_heading_and_fills_project_phrases():
    doc = Document()
    doc.add_paragraph("第三部分 响应文件格式")
    doc.add_paragraph("一、报价表")
    doc.add_paragraph("我方全面研究了“         ”采购文件（项目编号）后决定参加磋商。")
    doc.add_paragraph("本授权声明：（供应商名称）授权代表为我方 “       ” 项目（项目编号）的合法代理人。")

    fields = {
        "项目名称": "成都银行统一加密平台升级项目",
        "项目编号": "CYJC(X)-2025-0045",
        "供应商名称": "示例科技有限公司",
    }
    add_response_front_matter(doc, fields)
    fill_fields(doc, fields)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

    assert doc.paragraphs[0].text == "正本/副本"
    assert "第三部分 响应文件格式" not in "\n".join(paragraph.text for paragraph in doc.paragraphs[:12])
    assert "响 应 文 件" in text
    assert "项目名称：成都银行统一加密平台升级项目" in text
    assert "项目编号：CYJC(X)-2025-0045" in text
    assert "“成都银行统一加密平台升级项目”采购文件（项目编号：CYJC(X)-2025-0045）" in text
    assert "示例科技有限公司授权代表为我方 “成都银行统一加密平台升级项目” 项目（项目编号：CYJC(X)-2025-0045）" in text


def test_fill_fields_cleans_guarantee_html_and_adds_uppercase_amount():
    doc = Document()
    doc.add_paragraph("我方同意按照采购文件的要求，向贵单位交纳人民币    元（大写：        ）的磋商保证金。并承诺：下列任何情况发生时，我方将同意采购人没收我方参与磋商保证金：")
    raw = "参与磋商保证金 | 金额：人民币壹万柒仟元整（￥17,000.00）<br>缴纳方式：银行转账"

    fill_fields(doc, {"磋商保证金": raw})

    assert "人民币17,000.00元（大写：人民币壹万柒仟元整）" in doc.paragraphs[0].text
    assert "<br>" not in doc.paragraphs[0].text
    assert "缴纳方式" not in doc.paragraphs[0].text
    assert normalize_money_text(raw) == "17000.00"
    assert amount_to_rmb_upper("850000.00") == "人民币捌拾伍万元整"


def test_insert_sections_infers_heading_levels_from_numbered_lines():
    doc = Document()
    doc.add_paragraph("十、服务方案")
    sections = [
        {
            "章节编号": "10.1",
            "标题": "项目理解",
            "正文": "10.1.1 项目背景\n围绕统一加密平台升级项目进行响应。",
        }
    ]

    insert_sections(doc, sections)
    styled = [(paragraph.text, paragraph.style.name) for paragraph in doc.paragraphs if paragraph.text.strip()]

    assert ("10.1.1 项目背景", "Heading 3") in styled
    assert any(text == "围绕统一加密平台升级项目进行响应。" for text, _ in styled)


def test_cleanup_empty_headings_removes_blank_heading_paragraphs():
    doc = Document()
    doc.add_paragraph("十、服务方案")
    blank = doc.add_paragraph("")
    blank.style = "Heading 2"
    doc.add_paragraph("正文")

    removed = cleanup_empty_headings(doc)

    assert removed == 1
    assert all(not (paragraph.style.name == "Heading 2" and not paragraph.text.strip()) for paragraph in doc.paragraphs)


def test_update_existing_table_rows_preserves_row_count_and_fills_amounts():
    doc = Document()
    table = doc.add_table(rows=3, cols=6)
    headers = ["序号", "名称", "内容", "价款（含税）", "税率", "备注"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "产品一"
    table.rows[2].cells[0].text = "2"
    table.rows[2].cells[1].text = "产品二"

    fill_table_rows(
        table,
        [
            {"序号": "1", "价款（含税）": "10000", "税率": "6%"},
            {"序号": "2", "价款（含税）": "20000", "税率": "6%"},
        ],
        "update_existing_rows",
    )

    assert len(table.rows) == 3
    assert table.rows[1].cells[1].text == "产品一"
    assert table.rows[1].cells[3].text == "10000"
    assert table.rows[2].cells[3].text == "20000"
