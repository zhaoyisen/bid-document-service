from __future__ import annotations

import json
import re
import shutil
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from .generator import (
    OUTPUT_ROOT,
    file_response,
    generate_report,
    generate_xlsx,
    new_job_dir,
    safe_name,
    sanitize_generated_text,
    value,
)
from .models import GeneratedFile, GenerateResponse, TenderFormatFillRequest, BidPackageRequest


START_KEYWORDS = [
    "响应文件格式",
    "投标文件格式",
    "应答文件格式",
    "报价文件格式",
    "响应文件组成",
    "投标文件组成",
    "附件格式",
]

FORMAT_START_KEYWORDS = [
    "响应文件格式",
    "投标文件格式",
    "应答文件格式",
    "报价文件格式",
    "附件格式",
]

COMPOSITION_START_KEYWORDS = [
    "响应文件组成",
    "投标文件组成",
]

START_EXPLANATION_KEYWORDS = [
    "格式要求",
    "有关资格证明文件要求",
    "所有事项",
    "格式条款",
    "规范要求",
    "应认真阅读",
    "应包括",
    "包括下列",
    "按照采购文件要求",
    "作出实质性响应",
]

END_KEYWORDS = [
    "采购需求",
    "采购项目内容及要求",
    "采购内容及要求",
    "项目内容及要求",
    "技术要求",
    "合同条款",
    "合同格式",
    "评审办法",
    "评分办法",
    "用户需求书",
]

FIELD_LABELS = [
    "项目名称",
    "项目编号",
    "供应商名称",
    "投标人名称",
    "投标人",
    "报价总价",
    "报价",
    "税率",
    "法定代表人",
    "授权代表",
    "委托代理人",
    "联系人",
    "联系电话",
    "地址",
    "日期",
    "服务期",
    "交付期",
    "质保期",
]

BIDDER_FIELD_ALIASES = ("供应商名称", "投标人名称", "投标人", "响应供应商名称", "参与磋商供应商名称")


def local_name(element: Any) -> str:
    return element.tag.rsplit("}", 1)[-1]


def element_text(element: Any) -> str:
    texts = []
    for node in element.iter():
        if local_name(node) in {"t", "delText"} and node.text:
            texts.append(node.text)
    return "".join(texts).strip()


def start_keyword_present(text: str) -> bool:
    return any(keyword in text for keyword in START_KEYWORDS)


def normalize_heading_candidate(text: str) -> str:
    stripped = re.sub(r"\s+", " ", text.strip())
    stripped = re.sub(r"\t\s*\d+\s*$", "", stripped)
    stripped = re.sub(r"\.{2,}\s*\d+\s*$", "", stripped)
    stripped = re.sub(r"\s{2,}\d+\s*$", "", stripped)
    return stripped.strip()


def compact_text(text: str) -> str:
    return re.sub(r"\s+", "", text.strip())


def is_start_text(text: str) -> bool:
    candidate = normalize_heading_candidate(text)
    compact = compact_text(candidate)
    if not start_keyword_present(compact):
        return False
    if re.match(r"^[（(]\d+[）)]", compact):
        return False
    if any(keyword in compact for keyword in START_EXPLANATION_KEYWORDS):
        return False
    if len(compact) > 45:
        return False
    if re.match(r"^第[一二三四五六七八九十百千万0-9]+[部分章节篇]", compact):
        return True
    if re.match(r"^[一二三四五六七八九十百千万0-9]+[、.．]", compact):
        return True
    if re.match(r"^附件[一二三四五六七八九十百千万0-9]*[:：、.．]?", compact):
        return True
    if compact in START_KEYWORDS:
        return True
    return any(compact.startswith(keyword) and len(compact) <= len(keyword) + 8 for keyword in START_KEYWORDS)


def start_text_score(text: str) -> int:
    if not is_start_text(text):
        return 0
    compact = compact_text(normalize_heading_candidate(text))
    if any(keyword in compact for keyword in FORMAT_START_KEYWORDS):
        score = 100
    elif any(keyword in compact for keyword in COMPOSITION_START_KEYWORDS):
        score = 40
    else:
        score = 60
    if re.match(r"^第[一二三四五六七八九十百千万0-9]+[部分章节篇]", compact):
        score += 20
    return score


def is_toc_line(text: str) -> bool:
    stripped = text.strip()
    if "\t" in stripped and re.search(r"\t\s*\d+\s*$", stripped):
        return True
    if any(keyword in stripped for keyword in START_KEYWORDS + END_KEYWORDS) and re.search(r"\d+\s*$", stripped) and len(stripped) < 80:
        return True
    if re.search(r"\.{2,}\s*\d+\s*$", stripped):
        return True
    if re.search(r"\s{2,}\d+\s*$", stripped) and len(stripped) < 80:
        return True
    return False


def is_end_text(text: str) -> bool:
    if is_toc_line(text):
        return False
    if is_start_text(text):
        return False
    if not any(keyword in text for keyword in END_KEYWORDS):
        return False
    return bool(re.search(r"第[一二三四五六七八九十0-9]+[部分章节]", text) or re.match(r"^\d+[\.\s、]", text))


def find_format_range(doc: Document) -> tuple[int, int, list[str]]:
    warnings: list[str] = []
    body = doc._body._element
    children = list(body)
    start_idx = None
    start_candidates: list[tuple[int, int]] = []
    skipped_toc = False
    skipped_explanation = False
    for idx, child in enumerate(children):
        text = element_text(child)
        if start_keyword_present(text) and is_toc_line(text):
            skipped_toc = True
            continue
        score = start_text_score(text)
        if score:
            start_candidates.append((score, idx))
            continue
        if start_keyword_present(text):
            skipped_explanation = True
    if start_candidates:
        best_score = max(score for score, _ in start_candidates)
        start_idx = min(idx for score, idx in start_candidates if score == best_score)
    if start_idx is None:
        warnings.append("未识别到“响应文件格式/投标文件格式”等章节，已使用全文作为项目级模板。")
        return 0, len(children), warnings
    if skipped_toc:
        warnings.append("已跳过目录页中的响应文件格式条目，使用正文中的响应文件格式章节作为模板起点。")
    if skipped_explanation:
        warnings.append("已跳过说明性条款中的响应文件格式关键词，使用正式章节标题作为模板起点。")
    if any(idx < start_idx and score < best_score for score, idx in start_candidates):
        warnings.append("已跳过供应商须知中的响应文件组成说明，优先使用正式响应文件格式章节。")

    end_idx = len(children)
    for idx in range(start_idx + 1, len(children)):
        text = element_text(children[idx])
        if is_end_text(text):
            end_idx = idx
            break
    if end_idx == len(children):
        warnings.append("已识别响应文件格式起点，但未识别到下一章节，模板抽取到文档末尾。")
    return start_idx, end_idx, warnings


def keep_body_range(doc: Document, start_idx: int, end_idx: int) -> None:
    body = doc._body._element
    children = list(body)
    for idx, child in enumerate(children):
        if local_name(child) == "sectPr":
            continue
        if idx < start_idx or idx >= end_idx:
            body.remove(child)


def paragraph_text(paragraph: Paragraph) -> str:
    return "".join(run.text for run in paragraph.runs).strip()


def heading_level(paragraph: Paragraph, text: str) -> int | None:
    style_name = (paragraph.style.name or "") if paragraph.style else ""
    match = re.search(r"(\d+)$", style_name)
    if "Heading" in style_name and match:
        return int(match.group(1))
    m = re.match(r"^\s*(\d+(?:\.\d+){0,4})[\s、.．]", text)
    if m:
        return m.group(1).count(".") + 1
    if re.match(r"^[一二三四五六七八九十]+[、.．]", text):
        return 1
    return None


def extract_sections(doc: Document) -> list[dict[str, Any]]:
    sections = []
    for idx, paragraph in enumerate(doc.paragraphs, 1):
        text = paragraph_text(paragraph)
        if not text:
            continue
        level = heading_level(paragraph, text)
        if level:
            sections.append({
                "序号": len(sections) + 1,
                "层级": level,
                "标题": text,
                "段落序号": idx,
                "样式": paragraph.style.name if paragraph.style else "",
            })
    return sections


def table_text(cell) -> str:
    return "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())


def extract_tables(doc: Document) -> list[dict[str, Any]]:
    tables = []
    for idx, table in enumerate(doc.tables, 1):
        first_row = []
        if table.rows:
            first_row = [table_text(cell) for cell in table.rows[0].cells]
        tables.append({
            "表格序号": idx,
            "行数": len(table.rows),
            "列数": len(table.columns),
            "首行": first_row,
        })
    return tables


def get_table_headers(table) -> list[str]:
    if not table.rows:
        return []
    return [table_text(cell).replace("\n", " ").strip() for cell in table.rows[0].cells]


def detect_fields_in_text(text: str, location: str) -> list[dict[str, Any]]:
    fields = []
    for label in FIELD_LABELS:
        if label not in text:
            continue
        if re.search(rf"{re.escape(label)}\s*[:：]\s*($|[_\s　]+|/|年\s*月\s*日)", text):
            fields.append({
                "字段名": label,
                "位置": location,
                "原文": text[:180],
                "填充策略": "按冒号后的空白位置填充；无法确认时保留空白。",
            })
    return fields


def extract_field_mapping(doc: Document) -> list[dict[str, Any]]:
    fields: list[dict[str, Any]] = []
    seen = set()
    for idx, paragraph in enumerate(doc.paragraphs, 1):
        for item in detect_fields_in_text(paragraph_text(paragraph), f"段落{idx}"):
            key = (item["字段名"], item["位置"], item["原文"])
            if key not in seen:
                fields.append(item)
                seen.add(key)
    for table_idx, table in enumerate(doc.tables, 1):
        for row_idx, row in enumerate(table.rows, 1):
            for col_idx, cell in enumerate(row.cells, 1):
                for item in detect_fields_in_text(table_text(cell), f"表格{table_idx} 行{row_idx} 列{col_idx}"):
                    key = (item["字段名"], item["位置"], item["原文"])
                    if key not in seen:
                        fields.append(item)
                        seen.add(key)
    return fields


def write_json(path: Path, data: Any) -> Path:
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def extract_tender_format_file(source_docx: Path, project_name: str | None = None) -> GenerateResponse:
    source_docx = Path(source_docx)
    if source_docx.suffix.lower() != ".docx":
        raise ValueError("V2招标原格式抽取目前仅支持DOCX文件。PDF建议先转换为DOCX或作为解标来源。")

    name = project_name or source_docx.stem
    job_id, output_dir = new_job_dir(name)
    saved_source = output_dir / source_docx.name
    shutil.copy2(source_docx, saved_source)

    doc = Document(saved_source)
    start_idx, end_idx, warnings = find_format_range(doc)
    keep_body_range(doc, start_idx, end_idx)

    template_path = output_dir / f"{safe_name(name)}_项目级响应模板.docx"
    doc.save(template_path)

    template_doc = Document(template_path)
    sections = extract_sections(template_doc)
    tables = extract_tables(template_doc)
    fields = extract_field_mapping(template_doc)

    section_path = write_json(output_dir / f"{safe_name(name)}_章节清单.json", sections)
    table_path = write_json(output_dir / f"{safe_name(name)}_表格清单.json", tables)
    field_path = write_json(output_dir / f"{safe_name(name)}_字段映射.json", fields)
    report_path = output_dir / f"{safe_name(name)}_格式抽取报告.md"
    report_path.write_text(
        "\n".join([
            "# 招标原格式抽取报告",
            "",
            f"源文件：{source_docx.name}",
            f"模板范围：body元素 {start_idx} 到 {end_idx}",
            f"识别章节数：{len(sections)}",
            f"识别表格数：{len(tables)}",
            f"识别字段数：{len(fields)}",
            "",
            "## 提醒",
            *(f"- {warning}" for warning in warnings),
            "- 正式提交前仍需人工复核签字盖章、报价、授权、附件完整性和目录页码。",
        ]) + "\n",
        encoding="utf-8",
    )

    return GenerateResponse(
        job_id=job_id,
        files=[
            file_response(job_id, template_path, "docx-template"),
            file_response(job_id, field_path, "json"),
            file_response(job_id, section_path, "json"),
            file_response(job_id, table_path, "json"),
            file_response(job_id, report_path, "markdown"),
        ],
        warnings=warnings,
        metadata={
            "format_start_index": start_idx,
            "format_end_index": end_idx,
            "section_count": len(sections),
            "table_count": len(tables),
            "field_count": len(fields),
            "sections": sections[:80],
            "tables": tables[:80],
            "fields": fields[:120],
        },
    )


def iter_all_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def fill_label_text(text: str, label: str, fill_value: str) -> str:
    if not fill_value:
        return text
    seal_pattern = rf"({re.escape(label)}\s*[:：])\s*([（(][^）)]*[）)])\s*$"
    match = re.search(seal_pattern, text)
    if match:
        return text[:match.start()] + f"{match.group(1)}{fill_value}    {match.group(2)}" + text[match.end():]
    for pattern in [
        rf"({re.escape(label)}\s*[:：])\s*[_\s　/年月日.-]*$",
        rf"({re.escape(label)}\s*[:：])\s*[_\s　/.-]+",
    ]:
        match = re.search(pattern, text)
        if match:
            return text[:match.start()] + f"{match.group(1)}{fill_value}" + text[match.end():]
    return text


def expand_field_aliases(fields: dict[str, Any]) -> dict[str, str]:
    normalized = {str(k): "" if v is None else str(v) for k, v in fields.items()}
    bidder = (
        normalized.get("供应商名称")
        or normalized.get("投标人名称")
        or normalized.get("投标人")
        or normalized.get("响应供应商名称")
        or normalized.get("参与磋商供应商名称")
    )
    if bidder:
        for alias in BIDDER_FIELD_ALIASES:
            normalized.setdefault(alias, bidder)
    return normalized


def normalize_cell_label(text: str) -> str:
    text = re.sub(r"\s+", "", text)
    text = re.sub(r"[:：/（）()]", "", text)
    return text.strip()


def iter_distinct_row_cells(row):
    seen = set()
    for cell in row.cells:
        key = id(cell._tc)
        if key in seen:
            continue
        seen.add(key)
        yield cell


def fill_table_label_cells(doc: Document, fields: dict[str, str]) -> int:
    filled = 0
    normalized_labels = {normalize_cell_label(key): val for key, val in fields.items() if val}
    for table in doc.tables:
        for row in table.rows:
            cells = list(iter_distinct_row_cells(row))
            for idx, cell in enumerate(cells):
                label_text = normalize_cell_label(table_text(cell))
                if not label_text:
                    continue
                matched_value = ""
                for label, val in normalized_labels.items():
                    if label_text == label or (label_text.startswith(label) and len(label_text) <= len(label) + 4):
                        matched_value = val
                        break
                if not matched_value:
                    continue
                for target in cells[idx + 1:]:
                    if table_text(target).strip():
                        continue
                    set_cell(target, matched_value)
                    filled += 1
                    break
    return filled


def fill_fields(doc: Document, fields: dict[str, Any]) -> list[str]:
    warnings = []
    original_keys = {str(key) for key in fields}
    normalized = expand_field_aliases(fields)
    for paragraph in iter_all_paragraphs(doc):
        original = paragraph.text
        updated = original
        for key, val in normalized.items():
            updated = updated.replace(f"{{{{{key}}}}}", val)
            updated = updated.replace(f"{{{{ {key} }}}}", val)
            updated = fill_label_text(updated, key, val)
        if updated != original:
            paragraph.text = updated
    fill_table_label_cells(doc, normalized)
    doc_text = "\n".join(p.text for p in iter_all_paragraphs(doc))
    for key, val in normalized.items():
        if key not in original_keys:
            continue
        if key in BIDDER_FIELD_ALIASES and any(alias in doc_text for alias in BIDDER_FIELD_ALIASES):
            continue
        if val and key not in doc_text:
            warnings.append(f"未在模板中找到可直接填充的字段：{key}")
    fill_template_phrases(doc, normalized)
    return warnings


def insert_paragraph_after(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.text = sanitize_generated_text(text)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    return new_para


def insert_paragraph_before(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.text = sanitize_generated_text(text)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    return new_para


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def add_paragraph_safe(doc: Document, text: str, style: str | None = None) -> Paragraph:
    text = sanitize_generated_text(text)
    if not style:
        return doc.add_paragraph(text)
    try:
        return doc.add_paragraph(text, style=style)
    except KeyError:
        return doc.add_paragraph(text)


def add_heading_safe(doc: Document, text: str, level: int = 1) -> Paragraph:
    text = sanitize_generated_text(text)
    try:
        return doc.add_heading(text, level=level)
    except KeyError:
        return doc.add_paragraph(text)


def section_matches(text: str, section: dict[str, Any]) -> bool:
    title = value(section, "标题") or value(section, "title")
    number = value(section, "章节编号") or value(section, "number")
    if title and title in text:
        return True
    if number and text.strip().startswith(number):
        return True
    return False


def heading_style_name(level: int) -> str:
    return f"Heading {max(1, min(4, level))}"


def section_level(section: dict[str, Any], default: int = 2) -> int:
    raw = section.get("层级", section.get("level", default))
    try:
        return max(1, min(4, int(raw)))
    except Exception:
        return default


def section_text_lines(content: str) -> list[str]:
    lines = []
    for raw in sanitize_generated_text(content).splitlines():
        line = raw.strip()
        if not line or line in {"```", "```markdown"}:
            continue
        lines.append(line)
    return lines


def extract_json_payload(text: str) -> Any | None:
    clean = sanitize_generated_text(text)
    clean = re.sub(r"^```(?:json|markdown)?\s*", "", clean.strip(), flags=re.IGNORECASE)
    clean = re.sub(r"\s*```$", "", clean.strip())
    if not clean:
        return None
    candidates = [clean]
    start = clean.find("{")
    end = clean.rfind("}")
    if start != -1 and end > start:
        candidates.append(clean[start:end + 1])
    for candidate in candidates:
        try:
            return json.loads(candidate)
        except Exception:
            continue
    return None


def is_format_root_heading(text: str) -> bool:
    compact = compact_text(text)
    return bool(re.match(r"^第[一二三四五六七八九十百千万0-9]+[部分章节篇].{0,12}(响应文件格式|投标文件格式|应答文件格式)", compact))


def top_level_response_headings(doc: Document) -> list[str]:
    headings = []
    for paragraph in doc.paragraphs:
        text = paragraph_text(paragraph)
        if not text or is_format_root_heading(text):
            continue
        compact = compact_text(text)
        if re.match(r"^[一二三四五六七八九十百千万]+、", compact):
            headings.append(text)
        if len(headings) >= 18:
            break
    return headings


def add_response_front_matter(doc: Document, fields: dict[str, Any]) -> None:
    if not doc.paragraphs:
        return
    normalized = expand_field_aliases(fields)
    project = normalized.get("项目名称", "")
    project_code = normalized.get("项目编号", "")
    bidder = normalized.get("供应商名称") or normalized.get("投标人名称") or normalized.get("投标人") or ""
    headings = top_level_response_headings(doc)

    for paragraph in list(doc.paragraphs[:5]):
        if is_format_root_heading(paragraph_text(paragraph)):
            delete_paragraph(paragraph)
            break
    if not doc.paragraphs:
        return
    first = doc.paragraphs[0]
    lines = [
        "正本/副本",
        "",
        "响 应 文 件",
        "",
        f"项目名称：{project}",
        f"项目编号：{project_code}",
        f"参与磋商供应商名称：{bidder}",
        "",
        "日期：        年    月    日",
        "",
        "目 录",
        *headings,
        "",
    ]
    for line in lines:
        para = insert_paragraph_before(first, line)
        if line in {"正本/副本", "响 应 文 件", "目 录"}:
            try:
                para.alignment = 1
            except Exception:
                pass


def fill_template_phrases(doc: Document, fields: dict[str, Any]) -> None:
    normalized = expand_field_aliases(fields)
    project = normalized.get("项目名称", "")
    project_code = normalized.get("项目编号", "")
    bidder = normalized.get("供应商名称") or normalized.get("投标人名称") or normalized.get("投标人") or ""
    guarantee = (
        normalized.get("磋商保证金")
        or normalized.get("保证金金额")
        or normalized.get("保证金")
        or normalized.get("投标保证金")
    )
    for paragraph in iter_all_paragraphs(doc):
        text = paragraph.text
        updated = text
        if project:
            updated = re.sub(r"“\s*”(\s*采购文件)", f"“{project}”\\1", updated)
            updated = re.sub(r"“\s*”(\s*项目)", f"“{project}”\\1", updated)
        if project_code:
            updated = re.sub(r"（\s*项目编号\s*）", f"（项目编号：{project_code}）", updated)
            updated = re.sub(r"\(\s*项目编号\s*\)", f"（项目编号：{project_code}）", updated)
        if bidder:
            updated = re.sub(r"（\s*(参与磋商供应商的名称|参与磋商供应商名称|供应商名称|投标人名称|投标人)\s*）", bidder, updated)
            updated = re.sub(r"\(\s*(参与磋商供应商的名称|参与磋商供应商名称|供应商名称|投标人名称|投标人)\s*\)", bidder, updated)
        if guarantee:
            updated = re.sub(r"人民币\s*元（大写：\s*）", f"人民币{guarantee}元（大写：        ）", updated)
        if updated != text:
            paragraph.text = updated


def section_from_item(item: Any, idx: int) -> dict[str, Any] | None:
    if not isinstance(item, dict):
        return None
    title = item.get("标题") or item.get("title")
    content = item.get("正文") or item.get("content") or item.get("内容") or item.get("说明") or ""
    points = item.get("编制要点") or item.get("writing_points") or item.get("要点") or []
    if isinstance(points, str):
        points = [points]
    section = {
        "章节编号": str(item.get("章节编号") or item.get("number") or item.get("编号") or f"10.{idx}"),
        "层级": item.get("层级") or item.get("level") or 2,
        "标题": str(title or "未命名章节"),
        "正文": sanitize_generated_text(content),
        "编制要点": [sanitize_generated_text(point) for point in points if sanitize_generated_text(point)],
    }
    for key in ("关联要求编号", "所需证明材料", "缺口编号", "evidence_needed", "gap_ids"):
        if key in item:
            section[key] = item[key]
    return section


def sections_from_structured_text(text: str) -> list[dict[str, Any]]:
    parsed = extract_json_payload(text)
    if isinstance(parsed, dict):
        raw_sections = parsed.get("方案章节") or parsed.get("sections") or parsed.get("章节") or []
    elif isinstance(parsed, list):
        raw_sections = parsed
    else:
        raw_sections = []
    sections = []
    for idx, item in enumerate(raw_sections, 1):
        section = section_from_item(item, idx)
        if section:
            sections.append(section)
    return sections


def normalize_sections_for_insert(sections: list[dict[str, Any]]) -> list[dict[str, Any]]:
    normalized = []
    for idx, section in enumerate(sections, 1):
        if not isinstance(section, dict):
            parsed = sections_from_structured_text(str(section))
            normalized.extend(parsed)
            continue
        raw_nested = section.get("方案章节") or section.get("sections") or section.get("章节")
        if isinstance(raw_nested, list):
            for nested_idx, item in enumerate(raw_nested, 1):
                nested = section_from_item(item, nested_idx)
                if nested:
                    normalized.append(nested)
            continue
        content = value(section, "正文") or value(section, "content") or value(section, "内容")
        parsed = sections_from_structured_text(content)
        if parsed:
            normalized.extend(parsed)
            continue
        title = value(section, "标题") or value(section, "title")
        if not title and content:
            parsed = sections_from_structured_text(content)
            if parsed:
                normalized.extend(parsed)
                continue
        normalized.append(section_from_item(section, idx) or section)
    return normalized


def section_heading_text(section: dict[str, Any]) -> str:
    number = value(section, "章节编号") or value(section, "number")
    title = value(section, "标题") or value(section, "title")
    return f"{number} {title}".strip() or "未命名章节"


def section_insert_lines(section: dict[str, Any], include_heading: bool) -> list[tuple[str, str | None]]:
    lines: list[tuple[str, str | None]] = []
    if include_heading:
        lines.append((section_heading_text(section), heading_style_name(section_level(section))))
    content = value(section, "正文") or value(section, "content")
    for line in section_text_lines(content):
        lines.append((line, None))
    points = section.get("编制要点", section.get("writing_points", []))
    if isinstance(points, str):
        points = [points]
    for point in points:
        point_text = sanitize_generated_text(point)
        if point_text:
            lines.append((f"编制要点：{point_text}", "List Bullet"))
    evidence = value(section, "所需证明材料") or value(section, "evidence_needed")
    if evidence:
        lines.append((f"所需证明材料：{evidence}", None))
    gap = value(section, "缺口编号") or value(section, "gap_ids")
    if gap:
        lines.append((f"关联缺口：{gap}", None))
    return lines


def find_service_anchor(doc: Document) -> Paragraph | None:
    candidates = []
    for paragraph in doc.paragraphs:
        text = paragraph_text(paragraph)
        if not text:
            continue
        compact = compact_text(text)
        if "服务方案" in compact or "技术方案" in compact or "实施方案" in compact:
            candidates.append(paragraph)
    return candidates[-1] if candidates else None


def insert_sections(doc: Document, sections: list[dict[str, Any]]) -> list[str]:
    warnings = []
    sections = normalize_sections_for_insert(sections)
    paragraphs = list(doc.paragraphs)
    unmatched = []
    for section in sections:
        content = value(section, "正文") or value(section, "content")
        points = section.get("编制要点", section.get("writing_points", []))
        if isinstance(points, str):
            points = [points]
        matched = None
        for paragraph in paragraphs:
            if section_matches(paragraph.text, section):
                matched = paragraph
                break
        lines = []
        if content:
            lines.extend(section_text_lines(content))
        for point in points:
            point_text = sanitize_generated_text(point)
            if point_text:
                lines.append(f"编制要点：{point_text}")
        if not lines:
            continue
        if matched:
            for line in reversed(lines):
                insert_paragraph_after(matched, line)
        else:
            unmatched.append(section)
    if unmatched:
        service_anchor = find_service_anchor(doc)
        if service_anchor:
            insert_items: list[tuple[str, str | None]] = []
            for section in unmatched:
                insert_items.extend(section_insert_lines(section, include_heading=True))
            for text, style in reversed(insert_items):
                insert_paragraph_after(service_anchor, text, style=style)
            warnings.append(f"{len(unmatched)}个章节未匹配到具体子标题，已插入到招标原格式的“服务方案/技术方案”章节下。")
        else:
            add_heading_safe(doc, "自动补充方案内容", level=1)
            for section in unmatched:
                for text, style in section_insert_lines(section, include_heading=True):
                    add_paragraph_safe(doc, text, style=style)
            warnings.append(f"{len(unmatched)}个章节未匹配到招标原格式标题，已追加到文末“自动补充方案内容”。")
    return warnings


def match_table(doc: Document, table_index: int | None, match_headers: list[str]):
    if table_index is not None:
        idx = table_index - 1
        if 0 <= idx < len(doc.tables):
            return doc.tables[idx], f"表格{table_index}"
        return None, f"未找到序号为{table_index}的表格"
    expected = [h.strip() for h in match_headers if h.strip()]
    if not expected:
        return None, "未提供表格序号或表头关键词"
    for idx, table in enumerate(doc.tables, 1):
        header_text = " ".join(get_table_headers(table))
        if all(keyword in header_text for keyword in expected):
            return table, f"表格{idx}"
    return None, "未找到匹配表头关键词的表格：" + "、".join(expected)


def clear_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        paragraph.text = ""


def set_cell(cell, text: Any) -> None:
    clear_cell(cell)
    cell.text = sanitize_generated_text(text)


def delete_row(table, row_idx: int) -> None:
    row = table.rows[row_idx]
    table._tbl.remove(row._tr)


def value_for_header(row_data: dict[str, Any], header: str) -> str:
    candidates = [header, header.replace(" ", ""), header.split("\n")[0]]
    for key in candidates:
        if key in row_data:
            return sanitize_generated_text(row_data[key])
    for key, val in row_data.items():
        if key in header or header in key:
            return sanitize_generated_text(val)
    return ""


def row_number(row_data: dict[str, Any], fallback: int) -> str:
    raw = value_for_header(row_data, "序号") or value_for_header(row_data, "编号") or str(fallback)
    match = re.search(r"\d+", raw)
    return match.group(0) if match else raw.strip()


def update_existing_table_rows(table, rows: list[dict[str, Any]], headers: list[str]) -> None:
    existing_by_no: dict[str, int] = {}
    for idx in range(1, len(table.rows)):
        cells = table.rows[idx].cells
        if not cells:
            continue
        number = re.search(r"\d+", table_text(cells[0]).strip())
        if number:
            existing_by_no.setdefault(number.group(0), idx)

    for offset, row_data in enumerate(rows, 1):
        no = row_number(row_data, offset)
        row_idx = existing_by_no.get(no)
        if row_idx is None and offset < len(table.rows):
            row_idx = offset
        if row_idx is None:
            row = table.add_row()
        else:
            row = table.rows[row_idx]
        for idx, header in enumerate(headers):
            if idx >= len(row.cells):
                break
            cell_value = value_for_header(row_data, header)
            if cell_value:
                set_cell(row.cells[idx], cell_value)


def fill_table_rows(table, rows: list[dict[str, Any]], mode: str = "replace_data_rows") -> None:
    if not rows:
        return
    headers = get_table_headers(table)
    if not headers:
        headers = list(rows[0].keys())
        new_row = table.add_row()
        for idx, header in enumerate(headers):
            if idx < len(new_row.cells):
                set_cell(new_row.cells[idx], header)
    if mode == "replace_data_rows" and len(table.rows) > 1:
        for idx in range(len(table.rows) - 1, 0, -1):
            delete_row(table, idx)
    if mode == "update_existing_rows":
        update_existing_table_rows(table, rows, headers)
        return
    for row_data in rows:
        row = table.add_row()
        for idx, header in enumerate(headers):
            if idx >= len(row.cells):
                break
            cell_value = value_for_header(row_data, header)
            set_cell(row.cells[idx], cell_value)


def fill_tables(doc: Document, table_fills) -> list[str]:
    warnings = []
    for spec in table_fills:
        table, note = match_table(doc, spec.table_index, spec.match_headers)
        if table is None:
            warnings.append(note)
            continue
        fill_table_rows(table, spec.rows, spec.mode)
        warnings.append(f"{note}已填充{len(spec.rows)}行数据。")
    return warnings


def generate_from_tender_format(req: TenderFormatFillRequest) -> GenerateResponse:
    template_dir = OUTPUT_ROOT / safe_name(req.template_job_id)
    template_name = req.template_file_name or "项目级响应模板.docx"
    candidates = [template_dir / template_name]
    if req.template_file_name is None:
        candidates.extend(template_dir.glob("*项目级响应模板.docx"))
    template_path = next((p for p in candidates if p.exists()), None)
    if template_path is None:
        raise FileNotFoundError(f"未找到项目级模板：{template_dir} / {template_name}")

    job_id, output_dir = new_job_dir(req.project_name)
    doc = Document(template_path)
    fields = dict(req.fields)
    if req.project_name:
        fields.setdefault("项目名称", req.project_name)
    if req.bidder_name:
        fields.setdefault("供应商名称", req.bidder_name)
        fields.setdefault("投标人名称", req.bidder_name)
        fields.setdefault("投标人", req.bidder_name)

    add_response_front_matter(doc, fields)
    warnings = fill_fields(doc, fields)
    warnings.extend(fill_tables(doc, req.table_fills))
    warnings.extend(insert_sections(doc, req.sections))

    output_docx = output_dir / f"{safe_name(req.project_name)}_按招标原格式响应文件初稿.docx"
    doc.save(output_docx)

    package_req = BidPackageRequest(
        project_name=req.project_name,
        bidder_name=req.bidder_name,
        tender_name=req.tender_name,
        template_id=req.template_job_id,
        metadata={},
        response_matrix=req.response_matrix,
        material_gaps=req.material_gaps,
        sections=req.sections,
        checklist=req.checklist,
    )
    xlsx = generate_xlsx(package_req, output_dir)
    report = generate_report(package_req, output_dir, output_docx)

    return GenerateResponse(
        job_id=job_id,
        files=[
            file_response(job_id, output_docx, "docx"),
            file_response(job_id, xlsx, "xlsx"),
            file_response(job_id, report, "markdown"),
        ],
        warnings=warnings,
        metadata={
            "section_count": len(req.sections),
            "table_fill_count": len(req.table_fills),
            "response_matrix_count": len(req.response_matrix),
            "material_gap_count": len(req.material_gaps),
            "checklist_count": len(req.checklist),
        },
    )
