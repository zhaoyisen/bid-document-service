from __future__ import annotations

import hashlib
import hmac
import html
import re
import time
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any
from urllib.parse import quote

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from .checks import heading_depth, matrix_missing_fields, scan_docx_pollution
from .models import BidPackageRequest, GeneratedFile, GenerateResponse
from .schemas import CHECK_HEADERS, GAP_HEADERS, MATRIX_HEADERS
from .settings import get_settings


BASE_DIR = Path(__file__).resolve().parents[1]
OUTPUT_ROOT = get_settings().output_root

THINK_BLOCK_RE = re.compile(r"<think>[\s\S]*?</think>", re.IGNORECASE)
THINK_TAG_RE = re.compile(r"</?think>", re.IGNORECASE)


def now_string() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def safe_name(value: str, fallback: str = "bid") -> str:
    value = (value or fallback).strip()
    value = re.sub(r'[\\/:*?"<>|]+', "_", value)
    value = re.sub(r"\s+", "_", value)
    return value[:80] or fallback


def new_job_dir(project_name: str) -> tuple[str, Path]:
    job_id = f"{safe_name(project_name)}_{datetime.now().strftime('%Y%m%d%H%M%S')}_{uuid.uuid4().hex[:8]}"
    path = OUTPUT_ROOT / job_id
    path.mkdir(parents=True, exist_ok=True)
    return job_id, path


def download_signature(job_id: str, filename: str, expires: int, secret: str) -> str:
    message = f"{job_id}:{filename}:{expires}".encode("utf-8")
    return hmac.new(secret.encode("utf-8"), message, hashlib.sha256).hexdigest()


def verify_download_token(job_id: str, filename: str, expires: int, token: str) -> bool:
    settings = get_settings()
    secret = settings.download_secret
    if not secret:
        return False
    if expires < int(time.time()):
        return False
    expected = download_signature(job_id, filename, expires, secret)
    return hmac.compare_digest(expected, token or "")


def sanitize_generated_text(raw: Any) -> str:
    if raw is None:
        return ""
    text = str(raw)
    text = THINK_BLOCK_RE.sub("", text)
    text = re.sub(r"(?is)<think>.*", "", text)
    text = THINK_TAG_RE.sub("", text)
    text = re.sub(r"(?i)<br\s*/?>", "\n", text)
    text = re.sub(r"(?is)</p\s*>", "\n", text)
    text = re.sub(r"(?is)<[^>]+>", "", text)
    text = html.unescape(text)
    return text.strip()


def value(row: dict[str, Any], key: str) -> str:
    raw = row.get(key, "")
    if raw is None:
        return ""
    if isinstance(raw, (list, tuple)):
        return "；".join(sanitize_generated_text(item) for item in raw)
    return sanitize_generated_text(raw)


def collect_headers(rows: list[dict[str, Any]], preferred: list[str]) -> list[str]:
    headers = list(preferred)
    for row in rows:
        for key in row:
            if key not in headers:
                headers.append(key)
    return headers


def add_key_value_table(doc: Document, pairs: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, val in pairs:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = val or ""


def add_rows_table(doc: Document, rows: list[dict[str, Any]], headers: list[str], title: str, max_rows: int = 12) -> None:
    doc.add_heading(title, level=2)
    if not rows:
        doc.add_paragraph("暂无数据。")
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows[:max_rows]:
        cells = table.add_row().cells
        for idx, header in enumerate(headers):
            cells[idx].text = value(row, header)
    if len(rows) > max_rows:
        doc.add_paragraph(f"仅展示前 {max_rows} 条，完整数据见 Excel 工作底稿。")


def add_sections(doc: Document, sections: list[dict[str, Any]]) -> None:
    doc.add_heading("方案目录与正文要点", level=1)
    if not sections:
        doc.add_paragraph("暂无章节内容。")
        return
    for section in sections:
        number = value(section, "章节编号") or value(section, "number")
        title = value(section, "标题") or value(section, "title") or "未命名章节"
        level_raw = section.get("层级", section.get("level", 1))
        try:
            level = max(1, min(4, int(level_raw)))
        except Exception:
            level = 1
        heading = f"{number} {title}".strip()
        doc.add_heading(heading, level=level)
        content = value(section, "正文") or value(section, "content")
        if content:
            doc.add_paragraph(content)
        points = section.get("编制要点", section.get("writing_points", []))
        if isinstance(points, str):
            points = [points]
        if points:
            doc.add_paragraph("编制要点：")
            for point in points:
                point_text = sanitize_generated_text(point)
                if point_text:
                    doc.add_paragraph(point_text, style="List Bullet")
        evidence = value(section, "所需证明材料") or value(section, "evidence_needed")
        if evidence:
            doc.add_paragraph(f"所需证明材料：{evidence}")
        gap = value(section, "缺口编号") or value(section, "gap_ids")
        if gap:
            doc.add_paragraph(f"关联缺口：{gap}")


def generate_docx(req: BidPackageRequest, output_dir: Path) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{req.project_name}\n响应文件初稿")
    run.bold = True
    run.font.size = None

    doc.add_paragraph(f"生成时间：{now_string()}")
    doc.add_heading("项目基础信息", level=1)
    add_key_value_table(
        doc,
        [
            ("项目名称", req.project_name),
            ("投标人/供应商", req.bidder_name or ""),
            ("招标文件", req.tender_name or ""),
            ("模板ID", req.template_id or ""),
        ],
    )

    add_rows_table(doc, req.response_matrix, collect_headers(req.response_matrix, MATRIX_HEADERS[:8]), "响应矩阵摘要", max_rows=10)
    add_sections(doc, req.sections)
    add_rows_table(doc, req.material_gaps, collect_headers(req.material_gaps, GAP_HEADERS), "资料缺口清单", max_rows=20)
    add_rows_table(doc, req.checklist, collect_headers(req.checklist, CHECK_HEADERS), "提交前检查表", max_rows=20)

    doc.add_heading("交付说明", level=1)
    doc.add_paragraph("本文件为自动生成的投标响应初稿，正式提交前需复核报价、资质、案例、授权、签章、日期和附件完整性。")

    path = output_dir / f"{safe_name(req.project_name)}_响应文件初稿.docx"
    doc.save(path)
    return path


def format_sheet(ws) -> None:
    fill = PatternFill("solid", fgColor="D9EAF7")
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.fill = fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.freeze_panes = "A2"
    for column in ws.columns:
        letter = get_column_letter(column[0].column)
        max_len = max(len(str(cell.value or "")) for cell in column[:80])
        ws.column_dimensions[letter].width = min(max(max_len + 2, 12), 42)
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)


def append_sheet(wb: Workbook, title: str, rows: list[dict[str, Any]], preferred_headers: list[str]) -> None:
    ws = wb.create_sheet(title)
    headers = collect_headers(rows, preferred_headers)
    ws.append(headers)
    for row in rows:
        ws.append([value(row, header) for header in headers])
    format_sheet(ws)


def generate_xlsx(req: BidPackageRequest, output_dir: Path) -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = "项目说明"
    ws.append(["字段", "内容"])
    ws.append(["项目名称", req.project_name])
    ws.append(["投标人/供应商", req.bidder_name or ""])
    ws.append(["招标文件", req.tender_name or ""])
    ws.append(["模板ID", req.template_id or ""])
    ws.append(["生成时间", now_string()])
    format_sheet(ws)

    append_sheet(wb, "响应矩阵", req.response_matrix, MATRIX_HEADERS)
    append_sheet(wb, "资料缺口", req.material_gaps, GAP_HEADERS)
    append_sheet(wb, "方案目录", req.sections, ["章节编号", "层级", "标题", "编制要点", "所需证明材料", "缺口编号"])
    append_sheet(wb, "提交前检查表", req.checklist, CHECK_HEADERS)

    path = output_dir / f"{safe_name(req.project_name)}_投标工作底稿.xlsx"
    wb.save(path)
    return path


def generate_report(req: BidPackageRequest, output_dir: Path, docx_path: Path | None = None) -> Path:
    lines = [
        "# 提交前检查报告",
        "",
        f"项目名称：{req.project_name}",
        f"投标人/供应商：{req.bidder_name or ''}",
        f"生成时间：{now_string()}",
        "",
        "## 统计",
        f"- 响应矩阵条数：{len(req.response_matrix)}",
        f"- 资料缺口条数：{len(req.material_gaps)}",
        f"- 方案章节条数：{len(req.sections)}",
        f"- 检查项条数：{len(req.checklist)}",
        "",
        "## 高风险提示",
    ]
    high_risks = [
        row for row in req.response_matrix
        if value(row, "优先级") == "高" or value(row, "是否废标点") in {"是", "疑似"}
    ]
    if high_risks:
        for row in high_risks:
            lines.append(f"- {value(row, '编号')}: {value(row, '招标要求原文')[:120]} {value(row, '风险说明')}")
    else:
        lines.append("- 暂未识别到高风险项；仍需人工复核废标点、报价、签章和附件。")

    lines += ["", "## 资料缺口"]
    if req.material_gaps:
        for gap in req.material_gaps:
            lines.append(f"- {value(gap, '编号')}: {value(gap, '资料名称')} - {value(gap, '建议动作')}")
    else:
        lines.append("- 暂无资料缺口。")

    lines += ["", "## 提交前检查"]
    if req.checklist:
        for item in req.checklist:
            lines.append(f"- {value(item, '编号')}: {value(item, '检查项')} - {value(item, '状态') or '未完成'}")
    else:
        lines.append("- 暂无检查项。")

    lines += ["", "## 确定性检查"]
    missing = matrix_missing_fields(req.response_matrix)
    if missing:
        lines.append("- 响应矩阵缺少生产字段：" + "、".join(missing))
    else:
        lines.append("- 响应矩阵字段完整。")
    if docx_path and docx_path.exists():
        depth = heading_depth(docx_path)
        lines.append(f"- Word最大标题层级：{depth}")
        if depth < 3:
            lines.append("- 标题层级不足3级，服务/技术方案建议继续细化。")
        pollution = scan_docx_pollution(docx_path)
        if pollution:
            lines.append(f"- 正文残留风险：{len(pollution)}处")
            for item in pollution[:20]:
                lines.append(f"  - {item}")
        else:
            lines.append("- 未发现常见正文残留词。")

    path = output_dir / f"{safe_name(req.project_name)}_提交前检查报告.md"
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def file_response(job_id: str, path: Path, kind: str) -> GeneratedFile:
    settings = get_settings()
    safe_filename = quote(path.name)
    if settings.download_secret:
        expires = int(time.time()) + settings.public_download_ttl_seconds
        token = download_signature(job_id, path.name, expires, settings.download_secret)
        url = f"/public-files/{job_id}/{safe_filename}?expires={expires}&token={token}"
    else:
        url = f"/files/{job_id}/{safe_filename}"
    public_base_url = get_settings().public_base_url
    if public_base_url:
        url = f"{public_base_url}{url}"
    return GeneratedFile(
        name=path.name,
        type=kind,
        path=str(path),
        url=url,
    )


def generate_package(req: BidPackageRequest) -> GenerateResponse:
    warnings: list[str] = []
    job_id, output_dir = new_job_dir(req.project_name)
    docx = generate_docx(req, output_dir)
    xlsx = generate_xlsx(req, output_dir)
    report = generate_report(req, output_dir, docx)
    return GenerateResponse(
        job_id=job_id,
        files=[
            file_response(job_id, docx, "docx"),
            file_response(job_id, xlsx, "xlsx"),
            file_response(job_id, report, "markdown"),
        ],
        warnings=warnings,
    )
